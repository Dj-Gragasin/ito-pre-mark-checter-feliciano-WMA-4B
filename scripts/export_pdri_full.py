import csv
import json
import sys
from pathlib import Path

root = Path(r"c:\xampp\htdocs\cpionic")
docx_path = root / "PDRI-2018.docx"
out_json = root / "PDRI-2018.full.json"
out_csv = root / "PDRI-2018.full.csv"

try:
    import docx  # python-docx
except Exception:
    print("MISSING_PYTHON_DOCX")
    sys.exit(3)

doc = docx.Document(str(docx_path))

# Collect paragraphs for section hints
paragraphs = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]

def nearest_heading(before_idx: int) -> str:
    # best effort heading capture from recent non-empty lines
    candidates = []
    lo = max(0, before_idx - 12)
    for i in range(lo, before_idx + 1):
        t = doc.paragraphs[i].text.strip() if i < len(doc.paragraphs) else ""
        if t:
            candidates.append(t)
    for t in reversed(candidates):
        if any(k in t.lower() for k in ["recommended", "nutrient", "vitamin", "mineral", "energy", "intakes", "macronutrient"]):
            return t
    return candidates[-1] if candidates else ""

rows = []
max_cols = 0

# map table to nearest paragraph index by walking XML body order
body = doc._element.body
para_counter = 0
table_counter = 0
for el in body:
    tag = el.tag.lower()
    if tag.endswith('}p'):
        para_counter += 1
    elif tag.endswith('}tbl'):
        table_counter += 1
        heading = nearest_heading(max(0, para_counter - 1))
        table = doc.tables[table_counter - 1]
        for r_idx, row in enumerate(table.rows):
            cells = [c.text.replace('\n', ' ').strip() for c in row.cells]
            max_cols = max(max_cols, len(cells))
            rows.append({
                "section": heading,
                "table_index": table_counter,
                "row_index": r_idx,
                "cells": cells
            })

# JSON output
payload = {
    "source": str(docx_path.name),
    "table_count": table_counter,
    "row_count": len(rows),
    "rows": rows,
    "paragraphs": paragraphs,
}
out_json.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding='utf-8')

# CSV output (flattened)
headers = ["section", "table_index", "row_index"] + [f"col_{i+1}" for i in range(max_cols)]
with out_csv.open('w', newline='', encoding='utf-8-sig') as f:
    w = csv.writer(f)
    w.writerow(headers)
    for r in rows:
        base = [r["section"], r["table_index"], r["row_index"]]
        padded = r["cells"] + [""] * (max_cols - len(r["cells"]))
        w.writerow(base + padded)

print(f"OK|json={out_json}|csv={out_csv}|tables={table_counter}|rows={len(rows)}|max_cols={max_cols}|paragraphs={len(paragraphs)}")
